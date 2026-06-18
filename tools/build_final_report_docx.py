from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE


OUTPUT = "docs/BaoCao_CoffeeShopPOS_DesignPattern.docx"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(2)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(t.rows[0].cells[i], h, bold=True, color="FFFFFF")
        shade_cell(t.rows[0].cells[i], "7A4A32")
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
    if widths:
        for row in t.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()
    return t


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_code(doc, text):
    p = doc.add_paragraph()
    p.style = "Code"
    p.add_run(text)


def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BAO CAO DO AN CUOI KY")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(122, 74, 50)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MON: MAU THIET KE PHAN MEM")
    r.bold = True
    r.font.size = Pt(14)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DE TAI: PHAN MEM QUAN LY BAN HANG QUAN CA PHE\nCOFFEE SHOP POS")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(232, 132, 66)

    doc.add_paragraph()
    rows = [
        ("Ten ung dung", "PurrCoffee POS System"),
        ("Ngon ngu", "Java 17"),
        ("Giao dien", "Java Swing"),
        ("Kien truc", "Layered Architecture: Presentation - Service - Domain - Infrastructure"),
        ("Trong tam mon hoc", "Ap dung Design Pattern vao nghiep vu POS quan ca phe"),
        ("Trang thai code", "Da co demo chay duoc, test pass 18/18"),
    ]
    table(doc, ["Thong tin", "Noi dung"], rows, widths=[4.5, 11])
    doc.add_paragraph("Ghi chu: Tai lieu nay duoc tao de lam noi dung nen cho bao cao cuoi ky va gui nhom ve UML/Use Case trong Enterprise Architect.")
    doc.add_page_break()


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    for name, size, color in [
        ("Heading 1", 16, "7A4A32"),
        ("Heading 2", 13, "B86D2F"),
        ("Heading 3", 11.5, "4B3428"),
    ]:
        st = styles[name]
        st.font.name = "Arial"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(10)
        st.paragraph_format.space_after = Pt(6)

    if "Code" not in styles:
        code = styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code"]
    code.font.name = "Consolas"
    code.font.size = Pt(9)
    code.paragraph_format.left_indent = Cm(0.5)
    code.paragraph_format.space_after = Pt(4)
    return doc


def main():
    doc = setup_doc()
    add_title_page(doc)

    doc.add_heading("1. Gioi thieu de tai", level=1)
    doc.add_paragraph(
        "De tai Coffee Shop POS xay dung mot phan mem quan ly ban hang quan ca phe theo huong demo nghiep vu thuc te. "
        "He thong ho tro nhan vien thu ngan tao don, tuy bien thuc uong bang topping/size, ap dung khuyen mai, gui don sang bep, "
        "xu ly thanh toan qua cong Momo/VNPay gia lap va xem hoa don. Ben canh do, nhan vien pha che co man hinh xu ly don, "
        "quan tri vien co man hinh quan ly menu, topping, ton kho, nguoi dung, don hang va bao cao doanh thu."
    )
    doc.add_paragraph(
        "Trong tam cua do an khong phai la xay dung mot POS thuong mai hoan chinh, ma la chung minh cach ap dung cac mau thiet ke "
        "phan mem vao mot bai toan co nghiep vu ro rang. Code duoc to chuc theo kien truc phan lop de UI khong nam business logic."
    )

    doc.add_heading("2. Pham vi va muc tieu", level=1)
    add_bullets(doc, [
        "Xay dung ung dung Java Swing chay duoc bang JDK, khong phu thuoc Maven/Gradle.",
        "Phan quyen dang nhap theo vai tro: Admin, Cashier, Kitchen.",
        "Trien khai 7 Design Pattern: Decorator, Strategy, State, Observer, Factory Method, Singleton, Adapter.",
        "Cung cap luong demo day du: tao don -> tuy bien mon -> giam gia -> gui bep -> hoan tat -> thanh toan -> in/xem hoa don -> bao cao.",
        "Co test runner tu dong voi 18 test case de chung minh logic va pattern.",
        "Co script SQL thiet ke database, nhung runtime hien tai dung InMemoryRepository de demo nhanh va on dinh."
    ])

    doc.add_heading("3. Tong quan he thong", level=1)
    table(doc, ["Thanh phan", "Mo ta"], [
        ("LoginView", "Man hinh dang nhap, mo workspace theo role nguoi dung."),
        ("POSView", "Man hinh thu ngan: chon mon, tuy bien topping, quan ly gio hang, giam gia, gui bep, thanh toan."),
        ("KitchenView", "Man hinh bep: xem don dang cho/dang pha che, nhan don, hoan tat, huy don."),
        ("AdminView", "Man hinh quan tri: dashboard, menu, topping, don hang, lich su, ton kho, users, bao cao."),
        ("Services", "Dieu phoi use case: AuthService, OrderService, MenuService, PaymentService, ReportService, InventoryService, UserService."),
        ("Domain", "Chua entity va pattern classes: Order, OrderItem, Payment, Beverage, State/Strategy/Decorator/Observer/Factory/Adapter."),
        ("Infrastructure", "InMemoryRepository, MenuItemRecord, DatabaseConnection singleton demo.")
    ], widths=[4, 11])

    doc.add_heading("4. Kien truc phan lop", level=1)
    add_code(doc, "presentation -> service -> domain/infrastructure\nservice -> domain + infrastructure\ndomain khong phu thuoc presentation\ninfrastructure khong phu thuoc presentation")
    doc.add_paragraph(
        "Presentation chi xu ly giao dien va goi service. Service chiu trach nhiem dieu phoi nghiep vu. Domain chua mo hinh du lieu, "
        "luat trang thai va cac Design Pattern. Infrastructure cung cap repository va ket noi database demo."
    )
    table(doc, ["Tang", "Package/File tieu bieu", "Trach nhiem"], [
        ("presentation", "LoginView, POSView, KitchenView, AdminView, AppTheme", "Hien thi UI, bat su kien, goi service."),
        ("service", "OrderService, MenuService, PaymentService, ReportService, InventoryService, UserService", "Xu ly use case, validate, ket noi pattern voi nghiep vu."),
        ("domain", "Order, OrderItem, User, Payment, InventoryItem, domain.patterns.*", "Entity va pattern thuan OOP."),
        ("infrastructure", "InMemoryRepository, DatabaseConnection, MenuItemRecord", "Luu tru du lieu demo va schema ket noi DB.")
    ], widths=[3, 5, 7])

    doc.add_heading("5. Tac nhan va Use Case", level=1)
    doc.add_paragraph("Danh sach tac nhan nen ve trong Use Case Overview:")
    table(doc, ["Tac nhan", "Mo ta"], [
        ("Nhan vien thu ngan (Cashier)", "Dang nhap, tao don, chon mon, tuy bien mon, ap dung giam gia, gui bep, thanh toan, xem hoa don."),
        ("Nhan vien pha che (Kitchen)", "Dang nhap, xem don dang cho/dang pha che, nhan don, hoan tat don, huy don neu can."),
        ("Quan tri vien (Admin)", "Quan ly menu, topping, user, ton kho, xem don hang/lich su/bao cao."),
        ("Cong thanh toan (Momo/VNPay)", "Tac nhan ngoai xu ly thanh toan va tra transaction code."),
        ("He thong thong bao/bao cao", "Observer noi bo nhan su kien doi trang thai don de cap nhat man hinh/log.")
    ], widths=[4, 11])

    doc.add_heading("5.1 Use Case Overview de ve", level=2)
    table(doc, ["Actor", "Use case nen ve"], [
        ("Cashier", "Dang nhap, Tao don moi, Tim/loc menu, Chon thuc uong, Tuy bien topping/size, Them vao gio hang, Cap nhat so luong, Xoa mon, Ap dung khuyen mai, Gui don sang bep, Danh dau san sang, Thanh toan, Xem/Luu hoa don, Huy don."),
        ("Kitchen", "Dang nhap, Xem danh sach don, Nhan don, Hoan tat don, Huy don."),
        ("Admin", "Dang nhap, Xem dashboard, Quan ly menu, Quan ly topping, Quan ly ton kho, Quan ly nguoi dung, Xem don dang xu ly, Xem lich su don, Xem bao cao doanh thu/top mon."),
        ("Payment Gateway", "Xu ly thanh toan Momo, Xu ly thanh toan VNPay."),
        ("Observer/Logger", "Nhan thong bao thay doi trang thai don, Ghi log bao cao.")
    ], widths=[3.5, 11.5])

    doc.add_heading("5.2 Quan he include/extend goi y", level=2)
    table(doc, ["Use case chinh", "Include/Extend", "Use case lien quan"], [
        ("Tao don moi", "include", "Chon thuc uong, Them vao gio hang, Tinh tong tien"),
        ("Them vao gio hang", "extend", "Tuy bien topping/size"),
        ("Ap dung khuyen mai", "include", "Tinh lai tong tien bang DiscountStrategy"),
        ("Gui don sang bep", "include", "Kiem tra ton kho, Tru kho, Chuyen trang thai Pending -> Preparing"),
        ("Hoan tat don", "include", "Chuyen trang thai Preparing -> Ready, Thong bao thu ngan"),
        ("Thanh toan", "include", "Goi PaymentGateway Adapter, Luu Payment, Chuyen Ready -> Paid"),
        ("Xem bao cao", "include", "Thong ke doanh thu, Top mon ban chay"),
        ("Huy don", "extend", "Hoan tra ton kho neu don chua Paid")
    ], widths=[4.5, 2.5, 8])

    doc.add_heading("6. Dac ta Use Case chi tiet", level=1)
    use_cases = [
        ("UC01", "Dang nhap", "Admin/Cashier/Kitchen", "Nguoi dung co tai khoan active", "Nhap username/password -> AuthService xac thuc -> mo dung man hinh theo role", "Sai thong tin hoac user bi khoa thi hien thong bao loi"),
        ("UC02", "Tao don POS", "Cashier", "Cashier da dang nhap", "Tao Order Pending -> chon mon -> them item -> cap nhat subtotal/total", "Khong co item thi khong cho gui bep/thanh toan"),
        ("UC03", "Tuy bien thuc uong", "Cashier", "Da chon beverage", "Tick topping/size -> Decorator boc Beverage -> tinh description va price", "Topping khong active thi khong hien trong UI"),
        ("UC04", "Ap dung khuyen mai", "Cashier", "Order dang Pending/Ready va co item", "Chon discount -> OrderService gan DiscountStrategy -> recalculate", "Tong cuoi cung khong am"),
        ("UC05", "Gui don sang bep", "Cashier", "Order Pending va co item", "OrderService kiem tra kho -> tru kho -> PendingState.sendToKitchen -> Preparing", "Kho khong du thi throw InventoryException va don van Pending"),
        ("UC06", "Xu ly don tai bep", "Kitchen", "Co order active", "Kitchen chon order -> nhan/hoan tat -> Preparing -> Ready -> notify observers", "Trang thai khong hop le thi bi chan boi State Pattern"),
        ("UC07", "Thanh toan", "Cashier + Payment Gateway", "Order Ready", "Chon Momo/VNPay -> PaymentService goi adapter -> success thi Ready -> Paid va luu transaction code", "Gateway fail thi order van Ready, khong tao Payment"),
        ("UC08", "Quan ly menu/topping", "Admin", "Admin dang nhap", "Them/sua/disable beverage/topping qua MenuService", "Ten rong, gia am, category sai hoac topping trung bi reject"),
        ("UC09", "Quan ly ton kho/bao cao", "Admin", "Co du lieu don/kho", "Xem inventory, doanh thu, top mon ban chay", "Du lieu demo lay tu InMemoryRepository"),
        ("UC10", "Quan ly nguoi dung", "Admin", "Admin dang nhap", "Them user, khoa/mo khoa user, phan role", "Username trung hoac role sai bi reject")
    ]
    table(doc, ["Ma", "Ten use case", "Actor", "Tien dieu kien", "Luong chinh", "Ngoai le"], use_cases, widths=[1.4, 3, 2.5, 3, 4, 4])

    doc.add_heading("7. Ap dung Design Pattern", level=1)
    table(doc, ["Pattern", "Class chinh", "Ap dung trong he thong", "Test"], [
        ("Decorator", "Beverage, BeverageDecorator, PearlDecorator, LargeSizeDecorator, ExtraShotDecorator, MilkDecorator", "Tuy bien thuc uong bang topping/size ma khong tao class cho moi to hop", "TC02, TC15"),
        ("Strategy", "DiscountStrategy, NoDiscountStrategy, PercentDiscountStrategy, VipDiscountStrategy, BuyOneGetOneStrategy", "Thay doi thuat toan giam gia runtime", "TC03"),
        ("State", "OrderState, PendingState, PreparingState, ReadyState, PaidState, CancelledState", "Kiem soat vong doi don hang va chan thao tac sai trang thai", "TC04, TC05, TC15"),
        ("Observer", "OrderObserver, OrderEventPublisher, CashierScreen, KitchenScreen, ReportLogger", "Thong bao khi trang thai don thay doi, dac biet Ready de cap nhat cashier/kitchen/log", "TC07"),
        ("Factory Method", "BeverageFactory, CoffeeFactory, TeaFactory, MatchaFactory, SmoothieFactory", "Tao Beverage theo category trong MenuService/UI ma khong phu thuoc concrete class", "TC16"),
        ("Singleton", "AppConfig, DatabaseConnection", "Dam bao cau hinh app/ket noi DB demo co mot instance dung chung", "TC17"),
        ("Adapter", "PaymentGateway, PaymentResult, MomoAdapter, VnpayAdapter", "Chuan hoa cong thanh toan Momo/VNPay sau mot interface chung", "TC06, TC18")
    ], widths=[2.5, 4.5, 6, 2])

    doc.add_heading("8. Yeu cau ve Class Diagram", level=1)
    doc.add_paragraph("Nhom ve UML nen chia thanh 4 package dung voi code that:")
    add_bullets(doc, [
        "presentation: LoginView, POSView, KitchenView, AdminView, AppShell, AppTheme, PaymentDialog, ReceiptDialog.",
        "service: AuthService, OrderService, MenuService, PaymentService, ReportService, InventoryService, UserService, ReceiptService, ReceiptImageService.",
        "domain.model: User, Order, OrderItem, Payment, Topping, InventoryItem.",
        "domain.patterns.decorator/strategy/state/observer/factory/singleton/adapter.",
        "infrastructure: InMemoryRepository, DatabaseConnection, MenuItemRecord."
    ])
    doc.add_paragraph("Quan he quan trong can the hien:")
    add_bullets(doc, [
        "Order gom nhieu OrderItem va co Payment neu da thanh toan.",
        "Order co OrderState hien tai; cac state implement OrderState.",
        "OrderService su dung DiscountStrategy, OrderEventPublisher, InventoryService va Repository.",
        "PaymentService phu thuoc PaymentGateway interface, khong phu thuoc truc tiep Momo/VNPay concrete.",
        "BeverageDecorator implements Beverage va chua mot Beverage ben trong.",
        "BeverageFactory co cac concrete factory tao Beverage.",
        "KitchenView/observer classes nhan thong bao tu OrderEventPublisher.",
        "AppContext khoi tao repository, services va observers."
    ])

    doc.add_heading("9. Yeu cau ve Sequence Diagram", level=1)
    doc.add_paragraph("Nen ve toi thieu 5 sequence diagram sau:")
    table(doc, ["Sequence", "Doi tuong tham gia", "Thong diep chinh"], [
        ("Dang nhap theo role", "LoginView, AuthService, InMemoryRepository, User, POSView/KitchenView/AdminView", "login(username,password) -> findUser -> check password/status -> open view by role"),
        ("Tao don va them topping", "POSView, MenuService, BeverageFactory, BeverageDecorator, OrderService, Order", "select menu -> create beverage -> wrap decorators -> addItem -> recalculate"),
        ("Ap dung giam gia", "POSView, OrderService, DiscountStrategy, Order", "setDiscountStrategy -> calculateDiscount -> finalTotal = subtotal - discount"),
        ("Gui don sang bep", "POSView, OrderService, InventoryService, OrderState, OrderEventPublisher, KitchenView", "checkStock -> deduct -> sendToKitchen -> notify observers -> refresh kitchen"),
        ("Thanh toan Momo/VNPay", "POSView, PaymentDialog, PaymentService, PaymentGateway, PaymentResult, OrderState, ReceiptDialog", "processPayment -> success -> order.pay -> savePayment -> show receipt"),
        ("Admin quan ly menu", "AdminView, MenuService, BeverageFactory, InMemoryRepository", "add/update/disable beverage -> validate -> save/update item -> refresh list")
    ], widths=[3, 6, 6])

    doc.add_heading("10. Yeu cau ve State Diagram", level=1)
    doc.add_paragraph("Ve state diagram cho Order voi cac trang thai va chuyen trang thai sau:")
    add_code(doc, "Pending -> Preparing -> Ready -> Paid\nPending -> Cancelled\nPreparing -> Cancelled\nReady -> Cancelled\nPaid va Cancelled khoa moi thao tac sua/xu ly tiep")
    table(doc, ["Trang thai", "Cho phep", "Chan"], [
        ("Pending", "Them/xoa/cap nhat mon, gui bep, huy", "Thanh toan khi chua Ready"),
        ("Preparing", "Hoan tat sang Ready, huy neu chua Paid", "Sua mon"),
        ("Ready", "Thanh toan, huy neu can", "Gui bep lai"),
        ("Paid", "Khoa toan bo thao tac nghiep vu", "Sua, huy, gui bep"),
        ("Cancelled", "Khoa toan bo thao tac nghiep vu", "Sua, thanh toan, gui bep")
    ], widths=[3, 6, 6])

    doc.add_heading("11. Co so du lieu de ve ERD", level=1)
    doc.add_paragraph("Runtime dang dung InMemoryRepository de de demo, nhung file sql/schema.sql da mo ta cac bang chinh. ERD nen gom:")
    table(doc, ["Bang", "Thuoc tinh chinh", "Quan he"], [
        ("users", "id, username, password_hash, role, status", "role dung de phan quyen LoginView"),
        ("beverages", "id, name, base_price, category, active", "duoc chon trong POS/MenuService"),
        ("toppings", "id, name, extra_price, active", "duoc dung boi Decorator trong POS"),
        ("orders", "id, created_at, status, discount_type, subtotal, discount_amount, total_amount", "1 order co nhieu order_items va 0..1 payment"),
        ("order_items", "id, order_id, beverage_id, quantity, note, item_price", "thuoc mot order, lien ket beverage"),
        ("payments", "id, order_id, method, amount, transaction_code, status", "thanh toan thanh cong cua order"),
        ("inventory_items", "id, name, unit, quantity, reorder_level", "duoc InventoryService tru/rollback theo order")
    ], widths=[3, 5.5, 6.5])

    doc.add_heading("12. Kiem thu", level=1)
    table(doc, ["TC", "Noi dung", "Pattern/Module"], [
        ("TC01", "Dang nhap dung vao dung role", "Auth/Login"),
        ("TC02", "Ca phe sua + Tran chau + Size L tinh dung gia", "Decorator"),
        ("TC03", "Don 100.000d giam 10% con 90.000d", "Strategy"),
        ("TC04", "Pending -> Preparing -> Ready hop le", "State"),
        ("TC05", "Paid -> Preparing throw InvalidStateTransitionException", "State"),
        ("TC06", "Momo success thi order Paid va co transaction code", "Adapter/Payment"),
        ("TC07", "Order Ready thi Cashier observer nhan thong bao", "Observer"),
        ("TC08-TC09", "Admin CRUD menu va validate input", "MenuService/Admin"),
        ("TC10-TC11", "Tru kho, rollback, chan khi thieu kho", "Inventory/State"),
        ("TC12-TC13", "Receipt content va export PNG", "Receipt"),
        ("TC14", "Admin user add/lock validation", "UserService"),
        ("TC15", "Cart merge quantity va chan update sau khi gui bep", "OrderService/State"),
        ("TC16", "Factory tao beverage dung category", "Factory Method"),
        ("TC17", "Singleton tra ve cung instance", "Singleton"),
        ("TC18", "Payment gateway fail khong chuyen Paid", "Adapter failure path")
    ], widths=[2.2, 8, 5])
    doc.add_paragraph("Lenh chay test:")
    add_code(doc, "test.bat\nKet qua hien tai: All tests passed: 18/18")

    doc.add_heading("13. Kich ban demo ngan", level=1)
    add_numbered(doc, [
        "Dang nhap cashier01/123 de mo POS.",
        "Chon Ca phe sua, tick Tran chau va Size L, Add to cart de demo Decorator.",
        "Chon giam gia 10%, Apply discount de demo Strategy.",
        "Send kitchen de demo State Pending -> Preparing va tru kho.",
        "Dang nhap kitchen01/123, nhan/hoan tat don de demo Observer va State Ready.",
        "Quay lai cashier, thanh toan Momo/VNPay de demo Adapter va Payment.",
        "Dang nhap admin/123 de xem dashboard, menu, topping, inventory, users, report.",
        "Mo tai lieu pattern evidence neu giang vien hoi pattern nam o class nao."
    ])

    doc.add_heading("14. Han che va huong phat trien", level=1)
    add_bullets(doc, [
        "Runtime hien tai dung InMemoryRepository nen du lieu mat khi tat app; da co schema SQL de phat trien SQLite/MySQL sau.",
        "Khong dung Maven/JUnit de tranh loi dependency luc demo; TestRunner tu chay duoc bang JDK.",
        "Receipt hien co preview va PNG export, chua xuat PDF that.",
        "UI Swing da du demo, neu co thoi gian co the tiep tuc polish Admin/Kitchen hoac migrate JavaFX.",
        "Neu can demo nhieu may/dong bo realtime thi can persistent database hoac shared server."
    ])

    doc.add_heading("15. Checklist cho nhom ve so do", level=1)
    add_bullets(doc, [
        "Use Case Overview: actor Cashier, Kitchen, Admin, Payment Gateway; dung danh sach use case muc 5.",
        "Use Case chi tiet: ve rieng cho POS order, Kitchen order, Admin menu/user/report.",
        "Class Diagram: chia package 4 tang; nhan manh 7 pattern va cac service trung tam.",
        "Sequence Diagram: uu tien 5 flow o muc 9.",
        "State Diagram: ve vong doi Order o muc 10.",
        "ERD: ve bang users, beverages, toppings, orders, order_items, payments, inventory_items.",
        "Khi dat ten class trong so do, dung dung ten class trong tai lieu nay de khop code."
    ])

    doc.add_page_break()
    doc.add_heading("Phu luc A - Tai khoan demo", level=1)
    table(doc, ["Username", "Password", "Role", "Man hinh"], [
        ("admin", "123", "ADMIN", "AdminView"),
        ("cashier01", "123", "CASHIER", "POSView"),
        ("kitchen01", "123", "KITCHEN", "KitchenView"),
    ], widths=[3, 3, 3, 5])

    doc.add_heading("Phu luc B - File/tai lieu can nop kem", level=1)
    add_bullets(doc, [
        "Source code GitHub: PhucQuan/Coffee_Shop_POS_DesignPattern",
        "docs/PATTERN_EVIDENCE_TABLE.md",
        "docs/DEMO_SCRIPT.md",
        "docs/agent-plan/STATE.md",
        "sql/schema.sql",
        "run.bat, test.bat, generate-assets.bat"
    ])

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
